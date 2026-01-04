from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    """Add background color to table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_answer_key():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('AP Precalculus - Week 1 Classwork ANSWER KEY')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 73, 125)

    # Subtitle
    subtitle = doc.add_paragraph('Right Triangle Trig & Introduction to Radians')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.italics = True
    subtitle_run.font.size = Pt(12)

    # Instructions
    doc.add_paragraph()
    instructions = doc.add_paragraph('TEACHER INSTRUCTIONS: This key includes full solutions, common student errors, point values, and extension suggestions.')
    instructions.runs[0].font.bold = True
    instructions.runs[0].font.color.rgb = RGBColor(192, 0, 0)

    doc.add_paragraph()

    # PART A SOLUTIONS
    doc.add_heading('PART A: BASIC SKILLS - SOLUTIONS', level=1).runs[0].font.color.rgb = RGBColor(46, 117, 182)

    part_a_solutions = [
        {
            'num': 1,
            'problem': "In a right triangle, sin(θ) = 3/5. What is cos(θ)?",
            'solution': 'Using sin²(θ) + cos²(θ) = 1:\n(3/5)² + cos²(θ) = 1\n9/25 + cos²(θ) = 1\ncos²(θ) = 16/25\ncos(θ) = 4/5',
            'answer': '4/5 (or 0.8)',
            'common_error': 'Students might try to use SOH-CAH-TOA directly without remembering the Pythagorean identity.',
            'partial_credit': '2 pts if work shown, 1 pt if only answer given'
        },
        {
            'num': 2,
            'problem': "Convert 45° to radians.",
            'solution': '45° × (π/180°) = 45π/180 = π/4',
            'answer': 'π/4 radians',
            'common_error': 'Students might forget to reduce the fraction or forget to multiply by π.',
            'partial_credit': '2 pts for correct answer, 1 pt if answer is 45π/180 (unreduced)'
        },
        {
            'num': 3,
            'problem': "Convert π/6 radians to degrees.",
            'solution': '(π/6) × (180°/π) = 180/6 = 30°',
            'answer': '30°',
            'common_error': 'Students might forget the conversion factor or mix up which way to multiply.',
            'partial_credit': '2 pts for correct answer'
        },
        {
            'num': 4,
            'problem': "In a right triangle with opposite side = 7 and hypotenuse = 25, find sin(θ).",
            'solution': 'sin(θ) = opposite/hypotenuse = 7/25',
            'answer': '7/25 (or 0.28)',
            'common_error': 'Some students might incorrectly use adjacent instead of opposite.',
            'partial_credit': '2 pts for correct answer with units'
        },
        {
            'num': 5,
            'problem': "Convert 90° to radians.",
            'solution': '90° × (π/180°) = 90π/180 = π/2',
            'answer': 'π/2 radians',
            'common_error': 'Some students write "π/90" - they divided instead of multiplied.',
            'partial_credit': '2 pts for correct answer'
        },
        {
            'num': 6,
            'problem': "In a right triangle with adjacent side = 8 and hypotenuse = 17, find cos(θ).",
            'solution': 'cos(θ) = adjacent/hypotenuse = 8/17',
            'answer': '8/17 (or ≈0.47)',
            'common_error': 'Students might use opposite instead of adjacent, or confuse with sine.',
            'partial_credit': '2 pts for correct answer'
        },
        {
            'num': 7,
            'problem': "Convert 3π/4 radians to degrees.",
            'solution': '(3π/4) × (180°/π) = 3(180)/4 = 540/4 = 135°',
            'answer': '135°',
            'common_error': 'Students might only convert π to 180 and forget to multiply by 3/4.',
            'partial_credit': '1 pt if answer is 540°, 2 pts for correct'
        },
        {
            'num': 8,
            'problem': "If tan(θ) = 5/12, and this is a right triangle, what is sin(θ)?",
            'solution': 'tan(θ) = opposite/adjacent = 5/12\nUsing Pythagorean theorem: 5² + 12² = 25 + 144 = 169 = 13²\nHypotenuse = 13\nsin(θ) = opposite/hypotenuse = 5/13',
            'answer': '5/13 (or ≈0.385)',
            'common_error': 'Students might give tan(θ) value as answer, or forget to find hypotenuse.',
            'partial_credit': '1 pt if hypotenuse found (13), 2 pts for correct sin value'
        },
        {
            'num': 9,
            'problem': "Convert 120° to radians.",
            'solution': '120° × (π/180°) = 120π/180 = 2π/3',
            'answer': '2π/3 radians',
            'common_error': 'Unreduced answers like 120π/180 are acceptable with 1 pt penalty.',
            'partial_credit': '2 pts for reduced form, 1 pt for unreduced'
        },
        {
            'num': 10,
            'problem': "In a right triangle with opposite = 3 and adjacent = 4, find tan(θ).",
            'solution': 'tan(θ) = opposite/adjacent = 3/4',
            'answer': '3/4 (or 0.75)',
            'common_error': 'Students might write 4/3 (inverse).',
            'partial_credit': '2 pts for correct answer'
        },
        {
            'num': 11,
            'problem': "Convert 2π/3 radians to degrees.",
            'solution': '(2π/3) × (180°/π) = 2(180)/3 = 360/3 = 120°',
            'answer': '120°',
            'common_error': 'Some students compute 2/3 × π instead of the full conversion.',
            'partial_credit': '2 pts for correct answer'
        },
        {
            'num': 12,
            'problem': "What is sin(30°)?",
            'solution': 'From special triangle knowledge: sin(30°) = 1/2',
            'answer': '1/2 (or 0.5)',
            'common_error': 'Students might confuse with sin(60°) = √3/2.',
            'partial_credit': '2 pts for exact value, 1 pt for decimal approximation'
        },
        {
            'num': 13,
            'problem': "Convert 180° to radians.",
            'solution': '180° × (π/180°) = π',
            'answer': 'π radians',
            'common_error': 'Very common to forget the π symbol.',
            'partial_credit': '2 pts for π, 1 pt for "1 radian"'
        },
        {
            'num': 14,
            'problem': "What is cos(60°)?",
            'solution': 'From special triangle knowledge: cos(60°) = 1/2',
            'answer': '1/2 (or 0.5)',
            'common_error': 'Students might confuse with cos(30°) = √3/2.',
            'partial_credit': '2 pts for exact value'
        },
        {
            'num': 15,
            'problem': "Convert π radians to degrees.",
            'solution': 'π × (180°/π) = 180°',
            'answer': '180°',
            'common_error': 'Some students might write "π degrees".',
            'partial_credit': '2 pts for 180°'
        }
    ]

    for item in part_a_solutions:
        doc.add_paragraph(f"Problem {item['num']}: {item['problem']}", style='Heading 2')
        doc.add_paragraph(item['solution']).paragraph_format.left_indent = Inches(0.3)
        doc.add_paragraph(f"Answer: {item['answer']}", style='List Bullet')
        doc.add_paragraph(f"Common Error: {item['common_error']}", style='List Bullet').runs[0].font.italics = True
        doc.add_paragraph(f"Rubric: {item['partial_credit']}", style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # PART B SOLUTIONS
    doc.add_heading('PART B: INTERMEDIATE SKILLS - SOLUTIONS', level=1).runs[0].font.color.rgb = RGBColor(46, 117, 182)

    part_b_solutions = [
        {
            'num': 16,
            'problem': "A ladder leans against a wall, making a 50° angle with the ground. If the ladder is 20 feet long, how high does it reach on the wall?",
            'solution': 'The ladder is the hypotenuse = 20 ft. Height is the opposite side to the 50° angle.\nsin(50°) = opposite/hypotenuse = height/20\nheight = 20 × sin(50°) = 20 × 0.766 = 15.32 ft',
            'answer': '≈15.3 feet (or 20sin(50°))',
            'common_error': 'Students might use cos(50°) instead, or forget which side is opposite to the angle.',
            'extension': 'Ask: What if the angle was measured from the wall instead? How would that change the solution?'
        },
        {
            'num': 17,
            'problem': "An angle of 2π/3 radians sweeps an arc on a circle with radius 5. What is the arc length?",
            'solution': 'Arc length = θ × r (where θ is in radians)\nArc length = (2π/3) × 5 = 10π/3 units',
            'answer': '10π/3 units ≈ 10.47 units',
            'common_error': 'Students might forget to use the radius, or convert to degrees first (unnecessary).',
            'extension': 'Why is the formula simpler in radians than degrees?'
        },
        {
            'num': 18,
            'problem': "From the top of a 100-foot building, the angle of depression to a point on the ground is 25°. How far is that point from the base of the building?",
            'solution': 'tan(25°) = opposite/adjacent = 100/distance\ndistance = 100/tan(25°) = 100/0.466 ≈ 214.5 feet',
            'answer': '≈214.5 feet (or 100/tan(25°))',
            'common_error': 'Students confuse angle of depression with angle of elevation, or use wrong trig function.',
            'extension': 'Draw diagram showing angle of depression equals angle of elevation from the ground.'
        },
        {
            'num': 19,
            'problem': "A radian measure of π/6 corresponds to what degree measure? If this angle is on a circle with radius 12, what is the arc length?",
            'solution': 'Degrees: (π/6) × (180/π) = 30°\nArc length: (π/6) × 12 = 2π ≈ 6.28 units',
            'answer': '30°, Arc length = 2π ≈ 6.28 units',
            'common_error': 'Converting radians correctly is tricky - watch for multiplying vs dividing.',
            'extension': 'Verify: 30/360 × circumference = 30/360 × 2π(12) = 2π'
        },
        {
            'num': 20,
            'problem': "In a right triangle, one acute angle is 35° and the hypotenuse is 50 meters. Find both legs of the triangle.",
            'solution': 'sin(35°) = opposite/50, so opposite = 50sin(35°) ≈ 28.68 m\ncos(35°) = adjacent/50, so adjacent = 50cos(35°) ≈ 40.96 m',
            'answer': '≈28.7 m and ≈41.0 m (or 50sin(35°) and 50cos(35°))',
            'common_error': 'Using both functions correctly is key - students often mix them up.',
            'extension': 'Check with Pythagorean theorem: 28.68² + 40.96² ≈ 2500 = 50²'
        },
        {
            'num': 21,
            'problem': "A Ferris wheel has a radius of 30 feet. If a passenger sits at an angle of π/3 radians from the starting position, what arc length has the wheel rotated?",
            'solution': 'Arc length = θ × r = (π/3) × 30 = 10π feet',
            'answer': '10π feet ≈ 31.4 feet',
            'common_error': 'Forgetting to multiply by the radius.',
            'extension': 'If a full rotation is 2π radians, what fraction of the wheel has rotated?'
        },
        {
            'num': 22,
            'problem': "Convert: (a) 225° to radians, (b) 5π/6 radians to degrees. Then identify which quadrant each angle is in.",
            'solution': '(a) 225° × (π/180°) = 225π/180 = 5π/4 radians. Quadrant III (180°-270°)\n(b) (5π/6) × (180°/π) = 5(180)/6 = 150°. Quadrant II (90°-180°)',
            'answer': '(a) 5π/4, Quadrant III  (b) 150°, Quadrant II',
            'common_error': 'Quadrant identification - remember: Q1 (0-90), Q2 (90-180), Q3 (180-270), Q4 (270-360)',
            'extension': 'Draw these angles on the unit circle.'
        },
        {
            'num': 23,
            'problem': "A surveyor measures the angle of elevation to the top of a building as 40° from 150 feet away. How tall is the building?",
            'solution': 'tan(40°) = height/150\nheight = 150 × tan(40°) = 150 × 0.839 ≈ 125.85 feet',
            'answer': '≈125.9 feet (or 150tan(40°))',
            'common_error': 'Using sine or cosine instead of tangent.',
            'extension': 'What would the angle be if the surveyor was twice as far away?'
        },
        {
            'num': 24,
            'problem': "Given: cos(θ) = 3/5 in a right triangle. Find sin(θ) and tan(θ).",
            'solution': 'If cos(θ) = 3/5, then adjacent = 3, hypotenuse = 5\nUsing Pythagorean: 3² + opposite² = 5²\n9 + opposite² = 25\nopposite = 4\nsin(θ) = 4/5, tan(θ) = 4/3',
            'answer': 'sin(θ) = 4/5, tan(θ) = 4/3',
            'common_error': 'Forgetting to use Pythagorean theorem to find the third side.',
            'extension': 'Verify: sin²(θ) + cos²(θ) = (4/5)² + (3/5)² = 16/25 + 9/25 = 25/25 = 1'
        },
        {
            'num': 25,
            'problem': "A pendulum swings through an angle of 3π/8 radians. If the pendulum is 2 meters long, what distance does the bob travel?",
            'solution': 'Arc length (distance) = θ × r = (3π/8) × 2 = 6π/8 = 3π/4 meters',
            'answer': '3π/4 meters ≈ 2.36 meters',
            'common_error': 'Forgetting that the pendulum length is the radius.',
            'extension': 'What if it swings 3π/8 radians on each side of vertical? What is the total distance?'
        },
        {
            'num': 26,
            'problem': "Two angles measure 120° and 4π/3 radians. Are they equivalent? Explain.",
            'solution': 'Convert 4π/3 to degrees: (4π/3) × (180/π) = 4(180)/3 = 720/3 = 240°\n120° ≠ 240°, so NO they are not equivalent.\nThey differ by 120° (or 2π/3 radians).',
            'answer': 'No. 120° ≠ 240°',
            'common_error': 'Students might assume they are equivalent without converting.',
            'extension': 'What angle in degrees is equivalent to 4π/3? What about 4π/3 + 2π?'
        },
        {
            'num': 27,
            'problem': "A 16-foot ramp makes a 20° angle with the horizontal. What is the vertical rise and horizontal distance?",
            'solution': 'Vertical rise (opposite): sin(20°) = rise/16, so rise = 16sin(20°) ≈ 5.47 feet\nHorizontal distance (adjacent): cos(20°) = distance/16, so distance = 16cos(20°) ≈ 15.03 feet',
            'answer': 'Vertical rise ≈ 5.5 feet, Horizontal distance ≈ 15.0 feet',
            'common_error': 'Confusing which is opposite and adjacent relative to the angle given.',
            'extension': 'Check: 5.47² + 15.03² ≈ 256 = 16²'
        },
        {
            'num': 28,
            'problem': "If sin(θ) = 7/25, find cos(θ) and tan(θ) in the same right triangle.",
            'solution': 'If sin(θ) = 7/25, then opposite = 7, hypotenuse = 25\nUsing Pythagorean: 7² + adjacent² = 25²\n49 + adjacent² = 625\nadjacent² = 576\nadjacent = 24\ncos(θ) = 24/25, tan(θ) = 7/24',
            'answer': 'cos(θ) = 24/25, tan(θ) = 7/24',
            'common_error': 'Arithmetic errors in Pythagorean theorem - check 25² = 625.',
            'extension': 'This is the famous 7-24-25 Pythagorean triple!'
        },
        {
            'num': 29,
            'problem': "ACT Practice: A sector of a circle with radius 4 has a central angle of π/3. What is the perimeter of this sector?",
            'solution': 'Perimeter = arc length + 2 radii\nArc length = (π/3) × 4 = 4π/3\nPerimeter = 4π/3 + 4 + 4 = 4π/3 + 8',
            'answer': '4π/3 + 8 ≈ 12.19 units',
            'common_error': 'Forgetting to add the two radii, or calculating only arc length.',
            'extension': 'What if they ask for the area of the sector? (Formula: A = (1/2)r²θ = 8π/3)'
        },
        {
            'num': 30,
            'problem': "An angle of 1 radian sweeps an arc on a circle. If the arc length is 8 cm, what is the radius of the circle?",
            'solution': 'Arc length = θ × r\n8 = 1 × r\nr = 8 cm',
            'answer': '8 cm',
            'common_error': 'Students might try to convert 1 radian to degrees unnecessarily.',
            'extension': 'Why is 1 radian such a convenient unit for mathematics?'
        }
    ]

    for item in part_b_solutions:
        doc.add_paragraph(f"Problem {item['num']}: {item['problem']}", style='Heading 2')
        doc.add_paragraph(item['solution']).paragraph_format.left_indent = Inches(0.3)
        doc.add_paragraph(f"Answer: {item['answer']}", style='List Bullet')
        doc.add_paragraph(f"Common Error: {item['common_error']}", style='List Bullet').runs[0].font.italics = True
        doc.add_paragraph(f"Extension: {item['extension']}", style='List Bullet').runs[0].font.color.rgb = RGBColor(0, 128, 0)
        doc.add_paragraph()

    doc.add_page_break()

    # PART C SOLUTIONS
    doc.add_heading('PART C: CHALLENGE PROBLEMS - SOLUTIONS', level=1).runs[0].font.color.rgb = RGBColor(46, 117, 182)

    doc.add_paragraph('Note: Expect only 40-50% completion on these. These represent stretch goals and future topics.')
    doc.add_paragraph()

    part_c_solutions = [
        {
            'num': 31,
            'problem': "Prove: For any angle θ on the unit circle, sin²(θ) + cos²(θ) = 1.",
            'solution': 'On the unit circle, any point has coordinates (cos(θ), sin(θ)).\nBy definition of a circle: x² + y² = r²\nSince r = 1 on the unit circle: cos²(θ) + sin²(θ) = 1\nRearranging: sin²(θ) + cos²(θ) = 1 ✓'
        },
        {
            'num': 32,
            'problem': "Why is the radian a 'natural' unit for mathematics compared to degrees?",
            'solution': 'In radians, the arc length formula becomes simply s = θr (without any conversion factors).\nIn degrees, you would need: s = (θ/360) × 2πr, which is much more complicated.\nRadians eliminate the 2π and 360 from the formula, making calculus and higher math much cleaner.'
        },
        {
            'num': 33,
            'problem': "Find the exact value of sin(15°) using 15° = 45° - 30°.",
            'solution': 'This requires the angle difference formula (coming in Unit 3):\nsin(A - B) = sin(A)cos(B) - cos(A)sin(B)\nsin(45° - 30°) = sin(45°)cos(30°) - cos(45°)sin(30°)\n= (√2/2)(√3/2) - (√2/2)(1/2)\n= (√6/4) - (√2/4)\n= (√6 - √2)/4',
            'answer': '(√6 - √2)/4 ≈ 0.2588'
        },
        {
            'num': 34,
            'problem': "ACT Challenge: Circle with radius 6 cm, sector angle π/2.5 radians. Find: (a) degrees, (b) arc length, (c) sector area, (d) percentage of circle.",
            'solution': '(a) Degrees: (π/2.5) × (180/π) = 180/2.5 = 72°\n(b) Arc length: (π/2.5) × 6 = 6π/2.5 = 12π/5 ≈ 7.54 cm\n(c) Sector area: (1/2)r²θ = (1/2)(6²)(π/2.5) = 18π/2.5 = 36π/5 ≈ 22.62 cm²\n(d) Percentage: 72°/360° = 1/5 = 20%',
            'answer': '(a) 72°  (b) ≈7.54 cm  (c) ≈22.62 cm²  (d) 20%'
        },
        {
            'num': 35,
            'problem': "Extension: Cone with height 10 and base radius 4. Find the angle and arc length relationships.",
            'solution': 'Slant height: s = √(10² + 4²) = √(100 + 16) = √116 ≈ 10.77\nAngle with base: sin(α) = 10/√116, so α ≈ 67.4° ≈ 1.177 radians\nIf slant height wraps around base: arc length = 4 × 2π / (2π × 10.77 / 2πr) ... This is complex and shows how 3D problems relate to circular measurements.'
        }
    ]

    for item in part_c_solutions:
        doc.add_paragraph(f"Problem {item['num']}: {item['problem']}", style='Heading 2')
        doc.add_paragraph(item['solution']).paragraph_format.left_indent = Inches(0.3)
        if 'answer' in item:
            doc.add_paragraph(f"Answer: {item['answer']}", style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # Grading Rubric
    doc.add_heading('OVERALL GRADING RUBRIC', level=1)

    rubric_table = doc.add_table(rows=5, cols=2)
    rubric_table.style = 'Light Grid Accent 1'

    rubric_cells = [
        ('Category', 'Criteria'),
        ('Part A (40%)', 'Basic skills - Most problems are straightforward applications of formulas. Target 80%+ accuracy. 15 problems @ 2 pts each = 30 points'),
        ('Part B (40%)', 'Intermediate skills - Multi-step problems, real-world applications. Target 60-70% accuracy. 15 problems @ 3 pts each = 45 points'),
        ('Part C (20%)', 'Challenge problems - Extension and critical thinking. Target 40-50% accuracy acceptable. 5 problems @ 5 pts each = 25 points'),
        ('Total', '30 + 45 + 25 = 100 points')
    ]

    for i, (label, criteria) in enumerate(rubric_cells):
        row = rubric_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = criteria
        if i == 0:
            shade_cell(row.cells[0], 'D9E9F7')
            shade_cell(row.cells[1], 'D9E9F7')

    doc.add_paragraph()

    # Key Notes for Grading
    doc.add_heading('KEY NOTES FOR GRADING', level=2)
    doc.add_paragraph('Partial Credit: Always give partial credit for correct method with arithmetic errors.', style='List Bullet')
    doc.add_paragraph('Show Work: Require students to show work on Part B and Part C. Part A might be quick answers.', style='List Bullet')
    doc.add_paragraph('Unit Conversions: Accept radian answers that are unreduced (like 120π/180) with 1 point penalty.', style='List Bullet')
    doc.add_paragraph('Calculator Use: Parts B and C may use calculators. Part A emphasizes mental math and special angles.', style='List Bullet')
    doc.add_paragraph('Growth Mindset: Celebrate Part C attempts - this is where learning happens!', style='List Bullet')

    # Save document
    doc.save('/Users/franklyles/Documents/Claude Stuff/Disciplines/precal/week_1/Week_1_PreCal_Key_Classwork.docx')
    print("Answer key document created successfully!")

if __name__ == '__main__':
    create_answer_key()
