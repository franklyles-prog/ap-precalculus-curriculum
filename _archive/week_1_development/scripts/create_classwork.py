from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    """Add background color to table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_classwork():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('AP Precalculus Classwork - Week 1')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 73, 125)

    # Subtitle
    subtitle = doc.add_paragraph('Right Triangle Trig & Introduction to Radians')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.italics = True
    subtitle_run.font.size = Pt(12)

    # Dates
    dates = doc.add_paragraph('Sessions: January 7 & 9, 2026')
    dates.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dates_run = dates.runs[0]
    dates_run.font.italics = True
    dates_run.font.size = Pt(10)

    # Standards
    standards = doc.add_paragraph()
    standards_run = standards.add_run('Standards: PC.TR.1 (Radian measure), PC.TR.2 (Radian/Degree conversion), PC.TR.3 (Unit circle), PC.TR.4 (Special angles)')
    standards_run.font.italics = True
    standards_run.font.size = Pt(10)

    doc.add_paragraph()

    # Learning Objectives
    objectives_heading = doc.add_heading('Learning Objectives (SWBAT):', level=2)
    objectives_heading.runs[0].font.color.rgb = RGBColor(68, 114, 196)

    doc.add_paragraph('Apply SOH-CAH-TOA to solve right triangles', style='List Number')
    doc.add_paragraph('Convert angle measures between degrees and radians', style='List Number')
    doc.add_paragraph('Identify and evaluate special angles (30°, 45°, 60°) on the unit circle', style='List Number')
    doc.add_paragraph('Understand radians as arc length on the unit circle', style='List Number')

    doc.add_paragraph()

    # Reference Box
    doc.add_paragraph('FORMULA REFERENCE - Keep this handy!', style='Heading 2').runs[0].font.color.rgb = RGBColor(68, 114, 196)

    ref_table = doc.add_table(rows=1, cols=1)
    ref_table.autofit = False
    ref_cell = ref_table.rows[0].cells[0]
    shade_cell(ref_cell, 'D9E9F7')

    p = ref_cell.paragraphs[0]
    p.add_run('Right Triangle Trigonometry:\n').font.bold = True
    p.add_run('sin(θ) = opposite / hypotenuse      cos(θ) = adjacent / hypotenuse      tan(θ) = opposite / adjacent\n')
    p.add_run('(SOH-CAH-TOA)').font.italics = True
    p.add_run('\n\nDegree to Radian Conversion:\n').font.bold = True
    p.add_run('Radians = Degrees × (π/180)      Degrees = Radians × (180/π)')

    doc.add_paragraph()
    doc.add_paragraph()

    # PART A
    part_a_heading = doc.add_heading('PART A: BASIC SKILLS (Procedural Fluency) - 40%', level=1)
    part_a_heading.runs[0].font.color.rgb = RGBColor(46, 117, 182)

    instruction = doc.add_paragraph('These problems use direct application of SOH-CAH-TOA and basic conversions. Target: 80%+ accuracy')
    instruction.runs[0].font.italics = True
    instruction.runs[0].font.size = Pt(10)

    # Worked Example A
    example_heading = doc.add_paragraph('Worked Example A: Right Triangle Trig')
    example_heading.runs[0].font.bold = True
    example_heading.runs[0].font.color.rgb = RGBColor(46, 117, 182)

    example_table = doc.add_table(rows=1, cols=1)
    example_cell = example_table.rows[0].cells[0]
    shade_cell(example_cell, 'F2F2F2')

    ep = example_cell.paragraphs[0]
    ep.add_run('In a right triangle, the side opposite angle θ is 5 units, and the hypotenuse is 13 units. Find sin(θ), cos(θ), and tan(θ).\n\n')
    ep.add_run('Solution:\n').font.bold = True
    ep.add_run('Step 1: Identify the sides\n')
    ep.add_run('Opposite = 5, Adjacent = 12 (use Pythagorean theorem: 5² + 12² = 13²), Hypotenuse = 13\n\n')
    ep.add_run('Step 2: Apply SOH-CAH-TOA\n')
    ep.add_run('sin(θ) = opposite/hypotenuse = 5/13\n')
    ep.add_run('cos(θ) = adjacent/hypotenuse = 12/13\n')
    ep.add_run('tan(θ) = opposite/adjacent = 5/12\n\n')
    ep.add_run('Answer: sin(θ) = 5/13,  cos(θ) = 12/13,  tan(θ) = 5/12').font.bold = True
    ep.add_run('\n\nCommon Error to Avoid: ').font.color.rgb = RGBColor(192, 0, 0)
    ep.add_run("Don't confuse 'opposite' and 'adjacent' - always label relative to angle θ!").font.italics = True
    ep.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    # Part A Problems
    part_a_problems = [
        "In a right triangle, sin(θ) = 3/5. What is cos(θ)?",
        "Convert 45° to radians.",
        "Convert π/6 radians to degrees.",
        "In a right triangle with opposite side = 7 and hypotenuse = 25, find sin(θ).",
        "Convert 90° to radians.",
        "In a right triangle with adjacent side = 8 and hypotenuse = 17, find cos(θ).",
        "Convert 3π/4 radians to degrees.",
        "If tan(θ) = 5/12, and this is a right triangle, what is sin(θ)?",
        "Convert 120° to radians.",
        "In a right triangle with opposite = 3 and adjacent = 4, find tan(θ).",
        "Convert 2π/3 radians to degrees.",
        "What is sin(30°)?",
        "Convert 180° to radians.",
        "What is cos(60°)?",
        "Convert π radians to degrees."
    ]

    for i, problem in enumerate(part_a_problems, 1):
        p = doc.add_paragraph(problem, style='List Number')
        doc.add_paragraph('Answer: ___________________________').paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()

    # PART B
    part_b_heading = doc.add_heading('PART B: INTERMEDIATE SKILLS (Conceptual Understanding) - 40%', level=1)
    part_b_heading.runs[0].font.color.rgb = RGBColor(46, 117, 182)

    instruction_b = doc.add_paragraph('These problems require combining concepts and solving multi-step problems. Target: 60-70% accuracy')
    instruction_b.runs[0].font.italics = True
    instruction_b.runs[0].font.size = Pt(10)

    # Worked Example B
    example_b_heading = doc.add_paragraph('Worked Example B: Arc Length Application')
    example_b_heading.runs[0].font.bold = True
    example_b_heading.runs[0].font.color.rgb = RGBColor(46, 117, 182)

    example_b_table = doc.add_table(rows=1, cols=1)
    example_b_cell = example_b_table.rows[0].cells[0]
    shade_cell(example_b_cell, 'F2F2F2')

    ebp = example_b_cell.paragraphs[0]
    ebp.add_run('A point on the unit circle is located at angle π/4 radians. Find the coordinates (x, y) of this point and calculate the arc length traveled from (1, 0) to this point.\n\n')
    ebp.add_run('Solution:\n').font.bold = True
    ebp.add_run('Step 1: Recognize special angle\n')
    ebp.add_run('π/4 radians = 45° (special angle)\n\n')
    ebp.add_run('Step 2: Find coordinates using unit circle\n')
    ebp.add_run('At π/4: x = cos(π/4) = √2/2\n')
    ebp.add_run('At π/4: y = sin(π/4) = √2/2\n\n')
    ebp.add_run('Step 3: Calculate arc length\n')
    ebp.add_run('Arc length = θ × r (where r = 1 on unit circle)\n')
    ebp.add_run('Arc length = π/4 × 1 = π/4 units\n\n')
    ebp.add_run('Answer: Coordinates = (√2/2, √2/2),  Arc length = π/4 units').font.bold = True
    ebp.add_run('\n\nCommon Error to Avoid: ').font.color.rgb = RGBColor(192, 0, 0)
    ebp.add_run("Don't confuse radian measure with arc length - on unit circle, they're equal!").font.italics = True

    doc.add_paragraph()

    # Part B Problems
    part_b_problems = [
        "A ladder leans against a wall, making a 50° angle with the ground. If the ladder is 20 feet long, how high does it reach on the wall?",
        "An angle of 2π/3 radians sweeps an arc on a circle with radius 5. What is the arc length?",
        "From the top of a 100-foot building, the angle of depression to a point on the ground is 25°. How far is that point from the base of the building?",
        "A radian measure of π/6 corresponds to what degree measure? If this angle is on a circle with radius 12, what is the arc length?",
        "In a right triangle, one acute angle is 35° and the hypotenuse is 50 meters. Find both legs of the triangle.",
        "A Ferris wheel has a radius of 30 feet. If a passenger sits at an angle of π/3 radians from the starting position, what arc length has the wheel rotated?",
        "Convert the following angles: (a) 225° to radians, (b) 5π/6 radians to degrees. Then explain which quadrant each angle is in.",
        "A surveyor measures the angle of elevation to the top of a building as 40° from 150 feet away. How tall is the building?",
        "Given: cos(θ) = 3/5 in a right triangle. Find sin(θ) and tan(θ). (Hint: Use Pythagorean theorem)",
        "A pendulum swings through an angle of 3π/8 radians. If the pendulum is 2 meters long, what distance does the bob travel?",
        "Two angles measure 120° and 4π/3 radians. Are they equivalent? Explain why or why not.",
        "A 16-foot ramp makes a 20° angle with the horizontal. What is the vertical rise and horizontal distance?",
        "If sin(θ) = 7/25, find cos(θ) and tan(θ) in the same right triangle.",
        "ACT Practice: A sector of a circle with radius 4 has a central angle of π/3. What is the perimeter of this sector?",
        "An angle of 1 radian sweeps an arc on a circle. If the arc length is 8 cm, what is the radius of the circle?"
    ]

    for i, problem in enumerate(part_b_problems, 16):
        p = doc.add_paragraph(problem, style='List Number')
        doc.add_paragraph('Work:').paragraph_format.left_indent = Inches(0.5)
        doc.add_paragraph().paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()

    # PART C
    part_c_heading = doc.add_heading('PART C: CHALLENGE PROBLEMS (Critical Thinking & Extensions) - 20%', level=1)
    part_c_heading.runs[0].font.color.rgb = RGBColor(46, 117, 182)

    instruction_c = doc.add_paragraph('These problems require deeper thinking and connections to other concepts. Target: 40-50% accuracy is acceptable - this is growth!')
    instruction_c.runs[0].font.italics = True
    instruction_c.runs[0].font.size = Pt(10)

    # Part C Problems
    part_c_problems = [
        "Prove: For any angle θ on the unit circle, sin²(θ) + cos²(θ) = 1. (Hint: Use the Pythagorean theorem and the definition of sin/cos)",
        "Why is the radian a 'natural' unit for mathematics compared to degrees? Consider what happens in the formula Arc Length = θ × r.",
        "Find the exact value of sin(15°) by using the fact that 15° = 45° - 30°. (Hint: This will be important when we study angle difference formulas)",
        "ACT Challenge: A circle has radius 6 cm. An angle of π/2.5 radians creates a sector. Find: (a) the degree measure, (b) the arc length, (c) the sector area. Then find the percentage of the circle covered by this sector.",
        "Extension: Consider a 3D scenario. A cone has a height of 10 and a base radius of 4. The slant height creates an angle with the base. Find this angle in both radians and degrees, then find the arc length if the slant height were wrapped around the base circle."
    ]

    for i, problem in enumerate(part_c_problems, 31):
        p = doc.add_paragraph(problem, style='List Number')
        doc.add_paragraph('Work:').paragraph_format.left_indent = Inches(0.5)
        doc.add_paragraph().paragraph_format.left_indent = Inches(0.5)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("Remember: Show all your work for partial credit. Mathematics is about process, not just answers!")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.italics = True
    footer.runs[0].font.size = Pt(10)

    # Save document
    doc.save('/Users/franklyles/Documents/Claude Stuff/Disciplines/precal/week_1/Week_1_PreCal_Classwork.docx')
    print("Classwork document created successfully!")

if __name__ == '__main__':
    create_classwork()
